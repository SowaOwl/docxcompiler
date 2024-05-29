from docx import Document
from re import sub, search
from docx.table import Table
from typing import Dict, List
from docx.text.run import Run
from docx.oxml import parse_xml
from utils.insturctions import TYPES
from docx.text.paragraph import Paragraph
from docx.enum.table import WD_TABLE_ALIGNMENT
from utils.xmlStyles import border_table_style

def fill_data(doc: Document, data: Dict) -> None:
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, data)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, data)

    _fill_tables(doc, data)

def _replace_text_in_paragraph(paragraph: Paragraph, data: Dict) -> None:
    # Если в параграфе один run идет замена всего параграфа
    if len(paragraph.runs) == 1 : 
        original_text = paragraph.text
        replaced_text = _replace_text(paragraph.text, data)
        if replaced_text != original_text:
            paragraph.text = replaced_text
        
    else:
        repeat_count = 0
        while _check_braces(paragraph.runs, paragraph):
            if(repeat_count >= 10):
                raise Exception("Your document has incorrectly written elements or contains 10 or more variables in a paragraph")
            else:
                repeat_count += 1
            
            inside_brackets = False
            text_to_replace = ""
            runs_to_clear = []
            first_run_index = None

            for i, run in enumerate(paragraph.runs):
                original_text = run.text
                # Если run содержит в себе и открываюшие и закрываюшие скобки происходит замена текста всего run
                if "{{" in original_text and "}}" in original_text and not inside_brackets:
                    replaced_text = _replace_text(run.text, data)
                    run.text = replaced_text

                # Если в run только открываюшие то начинается отслеживание текста с начала {{ и добавляются runы для очстки
                elif "{{" in original_text and not inside_brackets:
                    start_index = original_text.find("{{")
                    text_to_replace = original_text[start_index:]
                    first_run_index = i
                    inside_brackets = True
                    runs_to_clear.append(run)

                # Если в run только закрываюшие то добавляется к тексту 
                # для замены окончание текста и происходит замена текста этот текст добовляется в run где была открываюшая скобка
                # а остальные зачишаются
                elif "}}" in original_text and inside_brackets:
                    end_index = original_text.find("}}") + 2
                    text_to_replace += original_text[:end_index]
                    replaced_text = _replace_text(text_to_replace, data)
                    paragraph.runs[first_run_index].text = replaced_text
                    inside_brackets = False
                    _clear_runs(runs_to_clear)
                    runs_to_clear = []

                    if end_index < len(original_text):
                        run.text = original_text[end_index:]
                    else:
                        run.text = ""

                else:
                    if inside_brackets:
                        text_to_replace += original_text
                        runs_to_clear.append(run)

            _clear_runs(runs_to_clear)

def _check_braces(runs: List[Run], paragraph: Paragraph) -> bool:
    pattern = r'\{\{' + TYPES['tables'] + r':.*:[^}]+}}'
    for run in runs:
        if search(pattern, paragraph.text):
            pass
        elif "{{" in run.text or "}}" in run.text:
            return True
    return False

def _clear_runs(runs_to_clear: List[Run]) -> None:
    for run_clear in runs_to_clear[1:]:
            run_clear.text = ""

def _fill_tables(doc: Document, data: Dict) -> None:
    if 'tables' in data:
        for item in data['tables']:
            pattern = r'\{\{' + TYPES['tables'] + r':' + item['name'] + r':[^}]+}}'
            for paragraph in doc.paragraphs:
                if search(pattern, paragraph.text):
                    _create_table(paragraph, item, doc)
                    paragraph.clear()

def _replace_text(text: str, data: Dict) -> str:
    old_text = text
    for type_name, items in data.items():
        for item in items:
            placeholder = r'\{\{' + TYPES[type_name] + r':' + item['name'] + r':[^}]+}}'
            match type_name:
                case 'stroke':
                    text = sub(placeholder, str(item['data']), text)
                    if(text == old_text):
                        placeholder = r'\{\{' + item['placeholder'] + r'\}\}'
                        text = sub(placeholder, str(item['data']), text)
                case 'logical':
                    logicalTemp = 'Да' if item['data'] else 'Нет'
                    text = sub(placeholder, str(item['placeholder']) + ' - ' + logicalTemp, text)
                case 'switcher':
                    logicalTemp = '✓' if item['data'] else '☓'
                    text = sub(placeholder, str(item['placeholder']) + ' - ' + logicalTemp, text)
                case 'tables':
                    pass
                case _:
                    text = sub(placeholder, str(item['data']), text)
    return text

def _create_table(paragraph: Paragraph, item: Dict, doc: Document) -> None:
    table = doc.add_table(rows=len(item['data']), cols=len(item['data'][0]))
    _move_table_after(table, paragraph)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True 
    _set_table_border(table)

    for i, row in enumerate(item['data']):
        for j, cell_txt in enumerate(row):
            table.cell(i, j).text = str(cell_txt)

def _move_table_after(table: Table, paragraph: Paragraph) -> None:
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

def _set_table_border(table: Table) -> None:
    table_element = table._element
    table_element.xpath('./w:tblPr')[0].append(parse_xml(border_table_style))