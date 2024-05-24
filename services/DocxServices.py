import re
import base64
import os
import subprocess
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from werkzeug.datastructures import FileStorage
from typing import Dict, List
from utils.insturctions import INSTRUCTIONS
from utils.types import TYPES
from utils.xmlStyles import border_table_style

class DocxServices:

    __TEMP_PATH = 'storage/'

    @staticmethod
    def extractFromDocx(file: FileStorage) -> dict:
        doc = Document(file)
        data = DocxServices._extractData(doc)
        return data
    
    @staticmethod
    def fillDataToFile(data: Dict) -> str:
        file_to_save_path = DocxServices.__TEMP_PATH + 'temp.docx'           
        doc = Document(DocxServices.__TEMP_PATH + 'test.docx')

        DocxServices._fillData(doc, data)
        doc.save(file_to_save_path)
        return DocxServices._getBase64AndDeleteFile(file_to_save_path)

    @staticmethod
    def _fillData(doc: Document, data: Dict) -> None:
        for paragraph in doc.paragraphs:
            DocxServices._replaceTextInParagraph(paragraph, data)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        DocxServices._replaceTextInParagraph(paragraph, data)

        DocxServices._fillTables(doc, data)

    @staticmethod
    def _replaceTextInParagraph(paragraph: Paragraph, data: Dict):
        def _clearRuns(runs_to_clear: List[Run]):
            for run_clear in runs_to_clear[1:]:
                    run_clear.text = ""

        def _checkBraces(runs: List[Run]):
            pattern = r'\{\{' + TYPES['tables'] + r':.*:[^}]+}}'
            for run in runs:
                if re.search(pattern, paragraph.text):
                    pass
                elif "{{" in run.text or "}}" in run.text:
                    return True
            return False

        if len(paragraph.runs) == 1 : 
            for run in paragraph.runs:
                original_text = run.text
                replaced_text = DocxServices._replaceText(run.text, data)
                if replaced_text != original_text:
                    run.text = replaced_text
        else:
            repeat_count = 0
            while _checkBraces(paragraph.runs):
                if(repeat_count >= 10):
                    raise Exception("Your document has incorrectly written elements or contains 10 or more variables in a paragraph")
                else:
                    repeat_count += 1
                
                inside_brackets = False
                text_to_replace = ""
                runs_to_clear = []
                first_run_index = None

                for i, run in enumerate(paragraph.runs):
                    original_text = run.text
                    if "{{" in original_text and "}}" in original_text and inside_brackets == False:
                        replaced_text = DocxServices._replaceText(run.text, data)
                        run.text = replaced_text

                    elif "{{" in original_text and inside_brackets == False:
                        start_index = original_text.find("{{")
                        text_to_replace = original_text[start_index:]
                        first_run_index = i
                        inside_brackets = True
                        runs_to_clear.append(run)

                    elif "}}" in original_text and inside_brackets:
                        end_index = original_text.find("}}") + 2
                        text_to_replace += original_text[:end_index]
                        replaced_text = DocxServices._replaceText(text_to_replace, data)
                        paragraph.runs[first_run_index].text = replaced_text
                        inside_brackets = False
                        _clearRuns(runs_to_clear)
                        runs_to_clear = []

                        if end_index < len(original_text):
                            run.text = original_text[end_index:]
                        else:
                            run.text = ""

                    else:
                        if inside_brackets:
                            text_to_replace += original_text
                            runs_to_clear.append(run)

                _clearRuns(runs_to_clear)

    @staticmethod
    def _getBase64AndDeleteFile(path: str) -> str:
        base64_string = ''
        pdf_path = DocxServices.__TEMP_PATH + 'temp.pdf'
        command = ['libreoffice', '--headless', '--convert-to', 'pdf', path, '--outdir', DocxServices.__TEMP_PATH]
        subprocess.run(command, check=True)
        
        with open(pdf_path, 'rb') as file:
            base64_string = base64.b64encode(file.read())
            
        os.remove(path)
        os.remove(pdf_path)
        
        return base64_string.decode('utf-8')

    @staticmethod
    def _fillTables(doc: Document, data: Dict) -> None:
        if 'tables' in data:
            for item in data['tables']:
                pattern = r'\{\{' + TYPES['tables'] + r':' + item['name'] + r':[^}]+}}'
                for paragraph in doc.paragraphs:
                    if re.search(pattern, paragraph.text):
                        DocxServices._createTable(paragraph, item, doc)
                        paragraph.clear()
    
    @staticmethod
    def _createTable(paragraph: Paragraph, item: Dict, doc) -> None:
        table = doc.add_table(rows=len(item['data']), cols=len(item['data'][0]))
        DocxServices._moveTableAfter(table, paragraph)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True 
        DocxServices._setTableBorder(table)

        for i, row in enumerate(item['data']):
            for j, cell_txt in enumerate(row):
                table.cell(i, j).text = str(cell_txt)

    @staticmethod
    def _moveTableAfter(table: Table, paragraph: Paragraph) -> None:
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    @staticmethod
    def _setTableBorder(table: Table) -> None:
        table_element = table._element
        table_element.xpath('./w:tblPr')[0].append(parse_xml(border_table_style))
    
    @staticmethod
    def _replaceText(text: str, data: Dict) -> str:
        test = text
        for type_name, items in data.items():
            for item in items:
                placeholder = r'\{\{' + TYPES[type_name] + r':' + item['name'] + r':[^}]+}}'
                match type_name:
                    case 'logical':
                        logicalTemp = 'Да' if item['data'] else 'Нет'
                        text = re.sub(placeholder, str(item['placeholder']) + ' - ' + logicalTemp, text)
                    case 'switcher':
                        logicalTemp = '✓' if item['data'] else '☓'
                        text = re.sub(placeholder, str(item['placeholder']) + ' - ' + logicalTemp, text)
                    case 'tables':
                        pass
                    case _:
                        text = re.sub(placeholder, str(item['data']), text)
        return text

    @staticmethod
    def _extractData(doc: Document) -> dict:
        data = {
            'stroke': [],
            'number': [],
            'logical': [],
            'switcher': [],
            'tables': []
        }
        
        for paragraph in doc.paragraphs:
            DocxServices._extractFromText(paragraph.text, data)
    
        for table in doc.tables:
            DocxServices._extractFromTable(table, data)
        
        doc.save(DocxServices.__TEMP_PATH + 'test.docx')

        return data
    
    @staticmethod
    def _extractFromText(text: str, data: Dict) -> None:
        for type_name, instr in INSTRUCTIONS.items():
            matches = re.findall(instr['pattern'], text)
            for match in matches:
                if(any(item['name'] == match[0] for item in data[type_name])):
                    continue
                else:
                    if len(instr['attrNames']) == 1:  # Если только одно имя атрибута
                        data[type_name].append({instr['attrNames'][0]: match})
                    else:                             # Если несколько имен атрибутов
                        data[type_name].append({name: value for name, value in zip(instr['attrNames'], match)})
    
    @staticmethod
    def _extractFromTable(table: Table, data: Dict) -> None:
        for row in table.rows:
            for cell in row.cells:
                DocxServices._extractFromText(cell.text, data)